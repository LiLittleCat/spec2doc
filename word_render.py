from __future__ import annotations

from pathlib import Path
import json
from typing import Optional
from copy import deepcopy

def render_word_docs(
    api_model: dict | None,
    db_model: dict | None,
    template_path: Optional[str],
    output_dir: str,
    include_api: bool = True,
    include_db: bool = True,
    output_filename: Optional[str] = None,
) -> list[str]:
    out_dir = Path(output_dir).resolve()
    out_files: list[str] = []
    api_model = api_model or {}
    db_model = db_model or {}

    if include_api and not include_db:
        out_path = out_dir / (output_filename or "Spec2Doc_接口设计文档.docx")
        _render_api_docx(api_model, template_path, str(out_path))
        out_files.append(str(out_path))
        return out_files

    if template_path:
        # 模板驱动：排版最好、最适合推广
        from docxtpl import DocxTemplate

        tpl = DocxTemplate(template_path)
        context = {
            "api": api_model,
            "db": db_model,
        }
        tpl.render(context)

        out_path = out_dir / (output_filename or "Spec2Doc_设计文档.docx")
        tpl.save(str(out_path))
        if include_api:
            _fill_endpoint_tables(str(out_path), api_model.get("endpoints", []))
        if include_db:
            _fill_db_tables(str(out_path), db_model.get("tables", []))
        out_files.append(str(out_path))
        return out_files

    # 无模板：先生成一个“简版文档”证明流程通了
    from docx import Document

    doc = Document()
    if include_db and not include_api:
        title = "数据库设计文档"
    else:
        title = f"{api_model.get('title', 'API')} 设计文档"
    doc.add_heading(title, level=0)

    # API section
    if include_api:
        doc.add_heading("接口设计", level=1)
        eps = api_model.get("endpoints", [])
        doc.add_paragraph(f"接口数量：{len(eps)}")

        for ep in eps[:200]:  # 防止太大先截断，可自行改
            doc.add_heading(f"{ep['method']} {ep['path']}", level=2)
            if ep.get("summary"):
                doc.add_paragraph(f"摘要：{ep['summary']}")
            if ep.get("operationId"):
                doc.add_paragraph(f"operationId：{ep['operationId']}")

    # DB section
    if include_db:
        doc.add_heading("数据库设计", level=1)
        tables = db_model.get("tables", [])
        doc.add_paragraph(f"表数量：{len(tables)}")

        for t in tables[:200]:
            doc.add_heading(f"表：{t['name']}", level=2)
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = "字段"
            hdr[1].text = "类型"
            hdr[2].text = "可空"
            hdr[3].text = "默认值"
            for c in t.get("columns", []):
                row = table.add_row().cells
                row[0].text = str(c.get("name", ""))
                row[1].text = str(c.get("type", ""))
                row[2].text = "YES" if c.get("nullable", True) else "NO"
                row[3].text = str(c.get("default", "") or "")

    out_path = out_dir / (output_filename or "Spec2Doc_简版.docx")
    doc.save(str(out_path))
    out_files.append(str(out_path))
    return out_files


def _format_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, str):
        return value
    try:
        return json.dumps(value, ensure_ascii=False, indent=2)
    except Exception:
        return str(value)


def _format_schema_hint(schema: dict) -> str:
    if not isinstance(schema, dict):
        return ""
    schema_type = schema.get("type", "")
    fmt = schema.get("format", "")
    enum = schema.get("enum", [])
    default = schema.get("default", "")
    parts = []
    if schema_type:
        parts.append(schema_type)
    if fmt:
        parts.append(f"format={fmt}")
    if enum:
        parts.append(f"enum={enum}")
    if default != "" and default is not None:
        parts.append(f"default={default}")
    return " ".join(parts)


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr[i].text = title
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = _format_value(value)


def _add_kv_table(doc, rows: list[tuple[str, object]]) -> None:
    data = [[k, _format_value(v)] for k, v in rows if _format_value(v)]
    _add_table(doc, ["字段", "值"], data)


def _format_security(security: list) -> str:
    if not security:
        return ""
    lines = []
    for item in security:
        if not isinstance(item, dict):
            continue
        parts = []
        for name, scopes in item.items():
            if scopes:
                parts.append(f"{name}({','.join(scopes)})")
            else:
                parts.append(str(name))
        if parts:
            lines.append(", ".join(parts))
    return "\n".join(lines)


def _format_server_variables(variables: dict) -> str:
    if not isinstance(variables, dict) or not variables:
        return ""
    lines = []
    for name, info in variables.items():
        if not isinstance(info, dict):
            lines.append(f"{name}={info}")
            continue
        default = info.get("default", "")
        enum = info.get("enum", [])
        desc = info.get("description", "")
        chunk = f"{name}={default}"
        if enum:
            chunk += f" enum={enum}"
        if desc:
            chunk += f" ({desc})"
        lines.append(chunk)
    return "\n".join(lines)


def _format_contact(contact: dict) -> str:
    if not isinstance(contact, dict):
        return ""
    parts = []
    for key in ("name", "email", "url"):
        value = contact.get(key, "")
        if value:
            parts.append(f"{key}={value}")
    return "；".join(parts)


def _format_license(license_obj: dict) -> str:
    if not isinstance(license_obj, dict):
        return ""
    name = license_obj.get("name", "")
    url = license_obj.get("url", "")
    if name and url:
        return f"{name} ({url})"
    return name or url


def _format_examples(examples: list[dict]) -> str:
    if not examples:
        return ""
    lines = []
    for ex in examples:
        if not isinstance(ex, dict):
            lines.append(_format_value(ex))
            continue
        name = ex.get("name", "")
        value = _format_value(ex.get("value"))
        ext = ex.get("externalValue", "")
        desc = ex.get("description", "")
        summary = ex.get("summary", "")
        label = name or summary or "example"
        detail = value
        if desc:
            detail = f"{detail}\n{desc}"
        if ext:
            detail = f"{detail}\nexternalValue: {ext}"
        lines.append(f"{label}: {detail}")
    return "\n\n".join(lines)


def _render_fields_table(doc, fields: list[dict], title: str) -> None:
    if not fields:
        return
    doc.add_paragraph(title)
    rows = []
    for f in fields:
        rows.append([
            f.get("name", ""),
            f.get("type", ""),
            "Y" if f.get("required") else "N",
            f.get("description", ""),
        ])
    _add_table(doc, ["字段", "类型", "必填", "描述"], rows)


def _render_nested_tables(doc, nested: dict, prefix: str) -> None:
    if not nested:
        return
    for path, info in nested.items():
        if not isinstance(info, dict):
            continue
        label = "数组项" if info.get("type") == "array" else "对象"
        title = f"{prefix}{path} ({label})"
        _render_fields_table(doc, info.get("fields", []), title)


def _render_response_headers(doc, headers: dict) -> None:
    if not isinstance(headers, dict) or not headers:
        return
    rows = []
    for name, header in headers.items():
        if isinstance(header, dict):
            schema = header.get("schema", {})
            rows.append([
                name,
                _format_schema_hint(schema),
                header.get("description", ""),
            ])
        else:
            rows.append([name, "", _format_value(header)])
    _add_table(doc, ["Header", "类型", "描述"], rows)


def _render_api_docx(api_model: dict, template_path: str | None, out_path: str) -> None:
    from docx import Document

    doc = Document(template_path) if template_path else Document()

    info = api_model.get("info", {}) or {}
    title = info.get("title") or api_model.get("title") or "API"
    version = info.get("version", "") or api_model.get("version", "")
    title_text = f"{title} 接口设计文档"
    if version:
        title_text = f"{title_text} v{version}"
    doc.add_heading(title_text, level=0)
    if info.get("description") or api_model.get("description"):
        doc.add_paragraph(info.get("description") or api_model.get("description"))

    doc.add_heading("基础信息", level=1)
    base_rows = [
        ("OpenAPI", api_model.get("openapi", "")),
        ("标题", info.get("title", "")),
        ("版本", info.get("version", "")),
        ("描述", info.get("description", "")),
        ("服务条款", info.get("termsOfService", "")),
        ("联系人", _format_contact(info.get("contact", {}))),
        ("许可证", _format_license(info.get("license", {}))),
    ]
    _add_kv_table(doc, base_rows)

    servers = api_model.get("servers", []) or []
    if servers:
        doc.add_heading("Servers", level=1)
        rows = []
        for server in servers:
            if not isinstance(server, dict):
                continue
            rows.append([
                server.get("url", ""),
                server.get("description", ""),
                _format_server_variables(server.get("variables", {})),
            ])
        _add_table(doc, ["URL", "描述", "变量"], rows)

    tags = api_model.get("tags", []) or []
    if tags:
        doc.add_heading("Tags", level=1)
        rows = []
        for tag in tags:
            if not isinstance(tag, dict):
                continue
            ext = tag.get("externalDocs", {})
            ext_text = ext.get("url", "") if isinstance(ext, dict) else ""
            rows.append([tag.get("name", ""), tag.get("description", ""), ext_text])
        _add_table(doc, ["名称", "描述", "外部文档"], rows)

    security = _format_security(api_model.get("security", []))
    if security:
        doc.add_heading("安全要求", level=1)
        doc.add_paragraph(security)

    extensions = api_model.get("extensions", {}) or {}
    if extensions:
        doc.add_heading("扩展字段", level=1)
        _add_kv_table(doc, list(extensions.items()))

    custom_rows = []
    standard_api_keys = {
        "openapi", "info", "title", "version", "description", "servers", "tags",
        "externalDocs", "security", "components", "extensions", "endpoints",
    }
    for key, value in api_model.items():
        if key not in standard_api_keys:
            custom_rows.append((key, value))
    if custom_rows:
        doc.add_heading("自定义字段", level=1)
        _add_kv_table(doc, custom_rows)

    endpoints = api_model.get("endpoints", []) or []
    doc.add_heading("接口清单", level=1)
    doc.add_paragraph(f"接口数量：{len(endpoints)}")

    for idx, ep in enumerate(endpoints, start=1):
        ep_title = f"{idx}. {ep.get('method', '')} {ep.get('path', '')}"
        doc.add_heading(ep_title, level=2)

        if ep.get("summary"):
            doc.add_paragraph(f"摘要：{ep.get('summary')}")
        if ep.get("description"):
            doc.add_paragraph(f"描述：{ep.get('description')}")

        info_rows = [
            ("方法", ep.get("method", "")),
            ("路径", ep.get("path", "")),
            ("OperationId", ep.get("operationId", "")),
            ("Tags", ", ".join(ep.get("tags", []) or [])),
            ("Deprecated", "是" if ep.get("deprecated") else "否"),
            ("路径摘要", ep.get("path_item_summary", "")),
            ("路径描述", ep.get("path_item_description", "")),
            ("Servers", _format_value(ep.get("servers", []))),
            ("安全要求", _format_security(ep.get("security", []))),
        ]
        _add_kv_table(doc, info_rows)

        param_fields = ep.get("param_fields", []) or []
        if param_fields:
            doc.add_heading("参数", level=3)
            rows = []
            for p in param_fields:
                example_text = _format_examples(p.get("examples", [])) or _format_value(p.get("example", ""))
                schema_hint = _format_schema_hint(p.get("schema", {}))
                type_text = p.get("type", "") or schema_hint
                rows.append([
                    p.get("name", ""),
                    p.get("in", ""),
                    type_text,
                    "Y" if p.get("required") else "N",
                    p.get("description", ""),
                    example_text,
                ])
            _add_table(doc, ["名称", "位置", "类型", "必填", "描述", "示例"], rows)

        request_body = ep.get("request_body", {}) or {}
        if request_body.get("content"):
            doc.add_heading("请求体", level=3)
            if request_body.get("description"):
                doc.add_paragraph(f"说明：{request_body.get('description')}")
            doc.add_paragraph(f"是否必填：{'是' if request_body.get('required') else '否'}")
            for item in request_body.get("content", []):
                if not isinstance(item, dict):
                    continue
                doc.add_paragraph(f"Content-Type：{item.get('content_type', '')}")
                fields = item.get("fields", []) or []
                _render_fields_table(doc, fields, "请求字段")
                _render_nested_tables(doc, item.get("nested", {}), "请求嵌套：")
                examples_text = _format_examples(item.get("examples", []))
                if examples_text:
                    doc.add_paragraph("请求示例：")
                    doc.add_paragraph(examples_text)

        responses_info = ep.get("responses_info", []) or []
        if responses_info:
            doc.add_heading("响应", level=3)
            for resp in responses_info:
                if not isinstance(resp, dict):
                    continue
                doc.add_heading(f"响应 {resp.get('status', '')}", level=4)
                if resp.get("description"):
                    doc.add_paragraph(f"描述：{resp.get('description')}")
                _render_response_headers(doc, resp.get("headers", {}))
                for item in resp.get("content", []) or []:
                    if not isinstance(item, dict):
                        continue
                    doc.add_paragraph(f"Content-Type：{item.get('content_type', '')}")
                    _render_fields_table(doc, item.get("fields", []), "响应字段")
                    _render_nested_tables(doc, item.get("nested", {}), "响应嵌套：")
                    examples_text = _format_examples(item.get("examples", []))
                    if examples_text:
                        doc.add_paragraph("响应示例：")
                        doc.add_paragraph(examples_text)

        ext_path = ep.get("path_item_extensions", {}) or {}
        ext_op = ep.get("extensions", {}) or {}
        if ext_path:
            doc.add_heading("路径扩展字段", level=3)
            _add_kv_table(doc, list(ext_path.items()))
        if ext_op:
            doc.add_heading("接口扩展字段", level=3)
            _add_kv_table(doc, list(ext_op.items()))

        callbacks = ep.get("callbacks", {})
        if callbacks:
            doc.add_heading("回调", level=3)
            doc.add_paragraph(_format_value(callbacks))

        ext_docs = ep.get("externalDocs", {})
        if ext_docs:
            doc.add_heading("外部文档", level=3)
            doc.add_paragraph(_format_value(ext_docs))

        custom_ep_rows = []
        standard_ep_keys = {
            "method", "path", "summary", "description", "operationId", "tags", "parameters",
            "requestBody", "responses", "security", "request_content_type", "req_fields",
            "req_nested", "resp_fields", "resp_nested", "param_fields", "request_body",
            "responses_info", "servers", "deprecated", "callbacks", "externalDocs",
            "path_item_summary", "path_item_description", "path_item_extensions",
            "extensions",
        }
        for key, value in ep.items():
            if key not in standard_ep_keys:
                custom_ep_rows.append((key, value))
        if custom_ep_rows:
            doc.add_heading("接口自定义字段", level=3)
            _add_kv_table(doc, custom_ep_rows)

    components = api_model.get("components", {}) or {}
    if components:
        doc.add_heading("组件", level=1)

        schemas = components.get("schemas", []) or []
        if schemas:
            doc.add_heading("Schemas", level=2)
            for schema in schemas:
                if not isinstance(schema, dict):
                    continue
                doc.add_heading(schema.get("name", ""), level=3)
                meta_rows = [
                    ("类型", schema.get("type", "")),
                    ("描述", schema.get("description", "")),
                ]
                _add_kv_table(doc, meta_rows)
                _render_fields_table(doc, schema.get("fields", []), "字段")
                _render_nested_tables(doc, schema.get("nested", {}), "嵌套：")

        comp_params = components.get("parameters", []) or []
        if comp_params:
            doc.add_heading("Parameters", level=2)
            rows = []
            for param in comp_params:
                if not isinstance(param, dict):
                    continue
                rows.append([
                    param.get("name", ""),
                    param.get("in", ""),
                    param.get("type", ""),
                    "Y" if param.get("required") else "N",
                    param.get("description", ""),
                ])
            _add_table(doc, ["名称", "位置", "类型", "必填", "描述"], rows)

        comp_responses = components.get("responses", []) or []
        if comp_responses:
            doc.add_heading("Responses", level=2)
            for resp in comp_responses:
                if not isinstance(resp, dict):
                    continue
                doc.add_heading(resp.get("name", ""), level=3)
                if resp.get("description"):
                    doc.add_paragraph(resp.get("description"))
                _render_response_headers(doc, resp.get("headers", {}))
                for item in resp.get("content", []) or []:
                    if not isinstance(item, dict):
                        continue
                    doc.add_paragraph(f"Content-Type：{item.get('content_type', '')}")
                    _render_fields_table(doc, item.get("fields", []), "响应字段")
                    _render_nested_tables(doc, item.get("nested", {}), "响应嵌套：")
                    examples_text = _format_examples(item.get("examples", []))
                    if examples_text:
                        doc.add_paragraph("响应示例：")
                        doc.add_paragraph(examples_text)

        comp_request_bodies = components.get("requestBodies", []) or []
        if comp_request_bodies:
            doc.add_heading("RequestBodies", level=2)
            for body in comp_request_bodies:
                if not isinstance(body, dict):
                    continue
                doc.add_heading(body.get("name", ""), level=3)
                if body.get("description"):
                    doc.add_paragraph(body.get("description"))
                doc.add_paragraph(f"是否必填：{'是' if body.get('required') else '否'}")
                for item in body.get("content", []) or []:
                    if not isinstance(item, dict):
                        continue
                    doc.add_paragraph(f"Content-Type：{item.get('content_type', '')}")
                    _render_fields_table(doc, item.get("fields", []), "请求字段")
                    _render_nested_tables(doc, item.get("nested", {}), "请求嵌套：")
                    examples_text = _format_examples(item.get("examples", []))
                    if examples_text:
                        doc.add_paragraph("请求示例：")
                        doc.add_paragraph(examples_text)

        comp_headers = components.get("headers", []) or []
        if comp_headers:
            doc.add_heading("Headers", level=2)
            rows = []
            for header in comp_headers:
                if not isinstance(header, dict):
                    continue
                rows.append([
                    header.get("name", ""),
                    header.get("type", ""),
                    header.get("description", ""),
                ])
            _add_table(doc, ["名称", "类型", "描述"], rows)

        comp_security = components.get("securitySchemes", []) or []
        if comp_security:
            doc.add_heading("SecuritySchemes", level=2)
            rows = []
            for scheme in comp_security:
                if not isinstance(scheme, dict):
                    continue
                rows.append([
                    scheme.get("name", ""),
                    scheme.get("type", ""),
                    scheme.get("scheme", "") or scheme.get("openIdConnectUrl", ""),
                    scheme.get("description", ""),
                ])
            _add_table(doc, ["名称", "类型", "方案/地址", "描述"], rows)

        comp_examples = components.get("examples", []) or []
        if comp_examples:
            doc.add_heading("Examples", level=2)
            for ex in comp_examples:
                if not isinstance(ex, dict):
                    continue
                doc.add_heading(ex.get("name", ""), level=3)
                if ex.get("summary"):
                    doc.add_paragraph(ex.get("summary"))
                if ex.get("description"):
                    doc.add_paragraph(ex.get("description"))
                doc.add_paragraph(_format_value(ex.get("value")))
                if ex.get("externalValue"):
                    doc.add_paragraph(f"externalValue: {ex.get('externalValue')}")

        comp_links = components.get("links", []) or []
        if comp_links:
            doc.add_heading("Links", level=2)
            rows = []
            for link in comp_links:
                if not isinstance(link, dict):
                    continue
                rows.append([
                    link.get("name", ""),
                    link.get("operationId", "") or link.get("operationRef", ""),
                    link.get("description", ""),
                ])
            _add_table(doc, ["名称", "目标", "描述"], rows)

        comp_callbacks = components.get("callbacks", []) or []
        if comp_callbacks:
            doc.add_heading("Callbacks", level=2)
            rows = []
            for cb in comp_callbacks:
                if not isinstance(cb, dict):
                    continue
                rows.append([
                    cb.get("name", ""),
                    ", ".join(cb.get("expressionKeys", []) or []),
                ])
            _add_table(doc, ["名称", "表达式键"], rows)

    doc.save(out_path)


def _fill_endpoint_tables(docx_path: str, endpoints: list[dict]) -> None:
    from docx import Document

    doc = Document(docx_path)
    tables = list(doc.tables)
    t_idx = 0
    
    # 使用集合避免重复，并记录表格 ID 而不是表格对象
    tables_to_remove_ids = set()
    
    for ep in endpoints:
        # 处理参数表格
        param_table, t_idx = _find_table_with_token(tables, "__PARAM_NAME__", t_idx)
        if param_table is not None:
            param_rows = _build_param_rows(ep)
            if not param_rows:
                tables_to_remove_ids.add(id(param_table))
            else:
                _fill_table_rows(param_table, "__PARAM_NAME__", param_rows)
        
        # 处理请求体表格
        req_table, t_idx = _find_table_with_token(tables, "__REQ_NAME__", t_idx)
        if req_table is not None:
            req_rows = _build_req_rows(ep)
            if not req_rows:
                tables_to_remove_ids.add(id(req_table))
            else:
                _fill_table_rows(req_table, "__REQ_NAME__", req_rows)
        
        # 处理请求体嵌套结构（使用 __REQ_NESTED_NAME__ 占位符）- 在响应表之前
        req_nested = ep.get("req_nested", {})
        req_nested_table, t_idx = _find_table_with_token(tables, "__REQ_NESTED_NAME__", t_idx)
        
        if req_nested_table is not None:
            if req_nested:
                from copy import deepcopy
                template_element = deepcopy(req_nested_table._element)
                parent = req_nested_table._element.getparent()
                
                nested_list = list(req_nested.items())
                first_path, first_info = nested_list[0]
                first_rows = _build_nested_rows(first_info, "__REQ_NESTED")
                if first_rows:
                    first_table_idx = list(parent).index(req_nested_table._element)
                    _insert_nested_label(doc, parent, first_table_idx, f"{first_path}", first_info)
                    _fill_table_rows(req_nested_table, "__REQ_NESTED_NAME__", first_rows)
                else:
                    tables_to_remove_ids.add(id(req_nested_table))
                
                # 后续嵌套：复制模板
                if len(nested_list) > 1:
                    first_table_idx = list(parent).index(req_nested_table._element)
                    insert_offset = 1
                    for nested_path, nested_info in nested_list[1:]:
                        _insert_nested_label(doc, parent, first_table_idx + insert_offset, f"{nested_path}", nested_info)
                        insert_offset += 1
                        cloned_table_elem = deepcopy(template_element)
                        parent.insert(first_table_idx + insert_offset, cloned_table_elem)
                        from docx.table import Table
                        cloned_table = Table(cloned_table_elem, doc)
                        nested_rows = _build_nested_rows(nested_info, "__REQ_NESTED")
                        if nested_rows:
                            _fill_table_rows(cloned_table, "__REQ_NESTED_NAME__", nested_rows)
                        else:
                            cloned_table._element.getparent().remove(cloned_table._element)
                        insert_offset += 1
            else:
                tables_to_remove_ids.add(id(req_nested_table))
        
        # 处理响应表格
        resp_table, t_idx = _find_table_with_token(tables, "__RESP_NAME__", t_idx)
        if resp_table is not None:
            resp_rows = _build_resp_rows(ep)
            if not resp_rows:
                tables_to_remove_ids.add(id(resp_table))
            else:
                _fill_table_rows(resp_table, "__RESP_NAME__", resp_rows)
        
        # 处理响应嵌套结构（使用 __RESP_NESTED_NAME__ 占位符）
        resp_nested = ep.get("resp_nested", {})
        resp_nested_table, t_idx = _find_table_with_token(tables, "__RESP_NESTED_NAME__", t_idx)
        
        if resp_nested_table is not None:
            if resp_nested:
                from copy import deepcopy
                template_element = deepcopy(resp_nested_table._element)
                parent = resp_nested_table._element.getparent()
                
                nested_list = list(resp_nested.items())
                first_path, first_info = nested_list[0]
                first_rows = _build_nested_rows(first_info, "__RESP_NESTED")
                if first_rows:
                    first_table_idx = list(parent).index(resp_nested_table._element)
                    _insert_nested_label(doc, parent, first_table_idx, f"{first_path}", first_info)
                    _fill_table_rows(resp_nested_table, "__RESP_NESTED_NAME__", first_rows)
                else:
                    tables_to_remove_ids.add(id(resp_nested_table))
                
                # 后续嵌套：复制模板
                if len(nested_list) > 1:
                    first_table_idx = list(parent).index(resp_nested_table._element)
                    insert_offset = 1
                    for nested_path, nested_info in nested_list[1:]:
                        _insert_nested_label(doc, parent, first_table_idx + insert_offset, f"{nested_path}", nested_info)
                        insert_offset += 1
                        cloned_table_elem = deepcopy(template_element)
                        parent.insert(first_table_idx + insert_offset, cloned_table_elem)
                        from docx.table import Table
                        cloned_table = Table(cloned_table_elem, doc)
                        nested_rows = _build_nested_rows(nested_info, "__RESP_NESTED")
                        if nested_rows:
                            _fill_table_rows(cloned_table, "__RESP_NESTED_NAME__", nested_rows)
                        else:
                            cloned_table._element.getparent().remove(cloned_table._element)
                        insert_offset += 1
            else:
                tables_to_remove_ids.add(id(resp_nested_table))
    
    # 删除所有空表 - 根据 ID 去重
    for table in tables:
        if id(table) in tables_to_remove_ids:
            try:
                table._element.getparent().remove(table._element)
            except Exception:
                # 表格可能已经被删除，忽略错误
                pass
    
    doc.save(docx_path)


def _insert_nested_label(doc, parent, insert_idx: int, nested_path: str, nested_info: dict) -> None:
    """在指定位置插入嵌套结构的标题段落"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    # 创建段落元素
    p = OxmlElement('w:p')
    
    # 创建段落属性
    pPr = OxmlElement('w:pPr')
    
    # 设置段前间距
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '200')  # 段前间距
    pPr.append(spacing)
    p.append(pPr)
    
    # 创建文本运行
    r = OxmlElement('w:r')
    
    # 设置加粗
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    
    # 设置字号
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')  # 10.5磅
    rPr.append(sz)
    r.append(rPr)
    
    # 添加文本
    t = OxmlElement('w:t')
    field_type = "数组项" if nested_info.get("type") == "array" else "对象"
    t.text = f"▸ {nested_path} ({field_type})"
    r.append(t)
    p.append(r)
    
    # 插入段落
    parent.insert(insert_idx, p)


def _find_table_with_token(tables: list, token: str, start_idx: int) -> tuple[Optional[object], int]:
    for i in range(start_idx, len(tables)):
        table = tables[i]
        for row in table.rows:
            for cell in row.cells:
                if token in cell.text:
                    return table, i + 1
    return None, len(tables)


def _fill_table_rows(table, token: str, rows: list[dict]) -> None:
    tpl_idx = None
    for i, row in enumerate(table.rows):
        if any(token in cell.text for cell in row.cells):
            tpl_idx = i
            break
    if tpl_idx is None:
        return

    tpl_tr = table.rows[tpl_idx]._tr
    if not rows:
        table._tbl.remove(tpl_tr)
        return

    for data in rows:
        new_tr = deepcopy(tpl_tr)
        table._tbl.append(new_tr)
        new_row = table.rows[-1]
        _apply_row_values(new_row, data)

    table._tbl.remove(tpl_tr)


def _apply_row_values(row, data: dict) -> None:
    for cell in row.cells:
        text = cell.text
        for key, value in data.items():
            if key in text:
                value_text = "" if value is None else str(value)
                if not value_text and key.endswith("_REQUIRED__"):
                    value_text = "N"
                cell.text = value_text


def _build_req_rows(ep: dict) -> list[dict]:
    rows = []
    for f in ep.get("req_fields", []) or []:
        rows.append({
            "__REQ_NAME__": f.get("name", ""),
            "__REQ_TYPE__": f.get("type", ""),
            "__REQ_REQUIRED__": "Y" if f.get("required") else "N",
            "__REQ_DESC__": f.get("description", ""),
        })
    return rows


def _build_resp_rows(ep: dict) -> list[dict]:
    rows = []
    for f in ep.get("resp_fields", []) or []:
        rows.append({
            "__RESP_NAME__": f.get("name", ""),
            "__RESP_TYPE__": f.get("type", ""),
            "__RESP_REQUIRED__": "Y" if f.get("required") else "N",
            "__RESP_DESC__": f.get("description", ""),
        })
    return rows


def _build_param_rows(ep: dict) -> list[dict]:
    """构建参数表格行（包括 query、path、header、cookie 等参数）"""
    rows = []
    for f in ep.get("param_fields", []) or []:
        rows.append({
            "__PARAM_NAME__": f.get("name", ""),
            "__PARAM_TYPE__": f.get("type", ""),
            "__PARAM_REQUIRED__": "Y" if f.get("required") else "N",
            "__PARAM_DESC__": f.get("description", ""),
        })
    return rows


def _build_nested_rows(nested_info: dict, prefix: str = "__NESTED") -> list[dict]:
    """构建嵌套结构表格行
    
    Args:
        nested_info: 嵌套结构信息
        prefix: 占位符前缀，如 "__REQ_NESTED" 或 "__RESP_NESTED"
    """
    rows = []
    for f in nested_info.get("fields", []) or []:
        rows.append({
            f"{prefix}_NAME__": f.get("name", ""),
            f"{prefix}_TYPE__": f.get("type", ""),
            f"{prefix}_REQUIRED__": "Y" if f.get("required") else "N",
            f"{prefix}_DESC__": f.get("description", ""),
        })
    return rows


def _fill_db_tables(docx_path: str, db_tables: list[dict]) -> None:
    """填充数据库表格模板
    
    模板中使用占位符:
    - __TABLE_NAME__: 表名
    - __TABLE_DESC__: 表描述/注释
    - __COL_NAME__: 列名
    - __COL_TYPE__: 列类型
    - __COL_REQUIRED__: 是否必填 (Y/N)
    - __COL_DESC__: 列描述
    """
    from docx import Document
    
    if not db_tables:
        return
    
    doc = Document(docx_path)
    
    # 找到包含 __TABLE_NAME__ 的段落作为模板
    template_para_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "__TABLE_NAME__" in para.text or "__TABLE_DESC__" in para.text:
            template_para_idx = i
            break
    
    # 找到包含 __COL_NAME__ 的表格作为模板
    template_table_idx = None
    template_table = None
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                if "__COL_NAME__" in cell.text:
                    template_table_idx = i
                    template_table = table
                    break
            if template_table is not None:
                break
        if template_table is not None:
            break
    
    if template_table is None:
        # 没有找到模板表格，跳过
        return
    
    # 获取模板段落（如果有）
    template_para = doc.paragraphs[template_para_idx] if template_para_idx is not None else None
    template_para_elem_tpl = deepcopy(template_para._element) if template_para else None

    template_table_elem = template_table._element
    template_table_elem_tpl = deepcopy(template_table_elem)
    
    # 处理第一个表
    first_table = db_tables[0]
    
    # 替换段落中的占位符
    if template_para:
        for run in template_para.runs:
            if "__TABLE_DESC__" in run.text:
                run.text = run.text.replace("__TABLE_DESC__", first_table.get("comment", "") or first_table.get("name", ""))
            if "__TABLE_NAME__" in run.text:
                run.text = run.text.replace("__TABLE_NAME__", first_table.get("name", ""))
    
    # 填充列数据
    col_rows = _build_col_rows(first_table)
    _fill_table_rows(template_table, "__COL_NAME__", col_rows)
    
    # 处理后续的表 - 需要复制模板段落和表格
    if len(db_tables) > 1:
        # 获取模板表格在文档中的位置
        table_parent = template_table_elem.getparent()
        template_table_idx_in_body = list(table_parent).index(template_table_elem)
        
        insert_idx = template_table_idx_in_body + 1
        
        for db_table in db_tables[1:]:
            # 复制段落
            if template_para_elem_tpl is not None:
                new_para_elem = deepcopy(template_para_elem_tpl)
                table_parent.insert(insert_idx, new_para_elem)
                
                # 替换新段落中的占位符
                from docx.text.paragraph import Paragraph
                new_para = Paragraph(new_para_elem, doc)
                for run in new_para.runs:
                    if "__TABLE_DESC__" in run.text:
                        run.text = run.text.replace("__TABLE_DESC__", db_table.get("comment", "") or db_table.get("name", ""))
                    if "__TABLE_NAME__" in run.text:
                        run.text = run.text.replace("__TABLE_NAME__", db_table.get("name", ""))
                insert_idx += 1
            
            # 复制表格
            new_table_elem = deepcopy(template_table_elem_tpl)
            table_parent.insert(insert_idx, new_table_elem)
            
            # 填充新表格
            from docx.table import Table
            new_table = Table(new_table_elem, doc)
            col_rows = _build_col_rows(db_table)
            _fill_table_rows(new_table, "__COL_NAME__", col_rows)
            insert_idx += 1
    
    doc.save(docx_path)


def _build_col_rows(db_table: dict) -> list[dict]:
    """构建列表格行数据"""
    rows = []
    for col in db_table.get("columns", []):
        # 确定是否必填：nullable 为 False 或者是主键
        nullable = col.get("nullable")
        required = "Y" if (nullable is False or col.get("is_pk", False)) else "N"
        rows.append({
            "__COL_NAME__": col.get("name", ""),
            "__COL_TYPE__": col.get("type", ""),
            "__COL_REQUIRED__": required,
            "__COL_DESC__": col.get("comment", "") or col.get("description", ""),
        })
    return rows
